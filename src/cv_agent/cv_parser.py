"""
CV Parser utility for extracting text from various file formats.
"""

import logging
from pathlib import Path
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class CVParser:
    """Utility class for parsing CV files of different formats"""

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from CV file based on its extension.

        Args:
            file_path: Path to the CV file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported or file cannot be read
        """
        path = Path(file_path)

        if not path.exists():
            raise ValueError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        try:
            if extension == ".pdf":
                return CVParser._extract_from_pdf(file_path)
            elif extension in [".docx", ".doc"]:
                return CVParser._extract_from_docx(file_path)
            elif extension == ".txt":
                return CVParser._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise ValueError(f"Failed to extract text from file: {str(e)}")

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        logger.info(f"Extracting text from PDF: {file_path}")

        text_content = []

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            full_text = "\n\n".join(text_content)

            if not full_text.strip():
                raise ValueError("PDF appears to be empty or contains only images")

            logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
            return full_text

        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise ValueError(f"Failed to read PDF file: {str(e)}")

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        logger.info(f"Extracting text from DOCX: {file_path}")

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            full_text = "\n\n".join(text_content)

            if not full_text.strip():
                raise ValueError("DOCX file appears to be empty")

            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise ValueError(f"Failed to read DOCX file: {str(e)}")

    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        logger.info(f"Extracting text from TXT: {file_path}")

        try:
            # Try different encodings
            encodings = ["utf-8", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()

                        if not text.strip():
                            raise ValueError("TXT file is empty")

                        logger.info(
                            f"Successfully extracted {len(text)} characters from TXT"
                        )
                        return text

                except UnicodeDecodeError:
                    continue

            raise ValueError("Failed to decode text file with supported encodings")

        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            raise ValueError(f"Failed to read TXT file: {str(e)}")
